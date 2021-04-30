#!/usr/bin/python

import sys
import json
from docx import Document
from google_trans_new import google_translator
translator= google_translator()
document = Document()



if len(sys.argv) == 2:

	file= sys.argv[1]
	with open(file, 'r') as f:
	    distros_dict = json.load(f)
	for distro in distros_dict:
	    	dos=distro['Vulnerabilities']
	document.add_heading('Criticidad Alta', level=1)
	for distro3 in dos:
		if distro3['Severity'] == "HIGH":
			document.add_paragraph(' ')
			p=document.add_paragraph('')
			name=p.add_run(distro3['PkgName'])
			name.bold = True
			name.italic = True
			p.add_run(' ('+ distro3['InstalledVersion']+ ') / ' + distro3['VulnerabilityID']).bold = True
			engl=document.add_paragraph('')
			j=engl.add_run(distro3['Description'])
			j.italic = True
			translation3 = translator.translate(distro3['Description'],lang_tgt='es')
			document.add_paragraph(translation3)
			document.add_paragraph('https://cve.mitre.org/cgi-bin/cvename.cgi?name='+ distro3['VulnerabilityID'])
			document.add_paragraph('')
			try:
				document.add_paragraph('Se arreglo en la version ' + distro3['FixedVersion'])
			except:
				exist=0
	document.add_heading('Criticidad Media', level=1)
	for distro2 in dos:
		if distro2['Severity'] == "MEDIUM":
			document.add_paragraph(' ')
			p2=document.add_paragraph('')
			name2=p2.add_run(distro2['PkgName'])
			name2.bold = True
			name2.italic = True
			p2.add_run(' ('+ distro2['InstalledVersion']+ ') / ' + distro2['VulnerabilityID']).bold = True
			engl2=document.add_paragraph('')
			j2=engl2.add_run(distro2['Description'])
			j2.italic = True
			translation2 = translator.translate(distro2['Description'],lang_tgt='es')
			document.add_paragraph(translation2)
			document.add_paragraph('https://cve.mitre.org/cgi-bin/cvename.cgi?name='+ distro2['VulnerabilityID'])
			document.add_paragraph('')
			try:
				document.add_paragraph('Se arreglo en la version ' + distro2['FixedVersion'])
			except:
				exist=0
	document.add_heading('Criticidad Baja', level=1)
	for distro1 in dos:
		if distro1['Severity'] == "LOW":
			document.add_paragraph(' ')
			p3=document.add_paragraph('')
			name3=p3.add_run(distro1['PkgName'])
			name3.bold = True
			name3.italic = True
			p3.add_run(' ('+ distro1['InstalledVersion']+ ') / ' + distro1['VulnerabilityID']).bold = True
			engl3=document.add_paragraph('')
			j2=engl3.add_run(distro1['Description'])
			j2.italic = True
			translation1 = translator.translate(distro1['Description'],lang_tgt='es')
			document.add_paragraph(translation1)
			document.add_paragraph('https://cve.mitre.org/cgi-bin/cvename.cgi?name='+ distro1['VulnerabilityID'])
			document.add_paragraph('')
			try:
				document.add_paragraph('Se arreglo en la version ' + distro3['FixedVersion'])
			except:
				exist=0
	document.save('demo.docx')
else:

	print('\n\033[91m[!]\033[0m \033[92mEl funcionamiento del script es: python3 ' + sys.argv[0] + ' archivo.json/archivo.txt \033[0m\n')
